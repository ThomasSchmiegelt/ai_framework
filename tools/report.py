"""
Engineering report generator: LaTeX → PDF (with python-docx fallback).
"""

import re
import subprocess
import tempfile
import uuid
from pathlib import Path

REPORTS_DIR = Path("data/reports")
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


def _latex_escape(text: str) -> str:
    replacements = [
        ("\\", r"\textbackslash{}"),
        ("&", r"\&"), ("%", r"\%"), ("$", r"\$"),
        ("#", r"\#"), ("_", r"\_"), ("{", r"\{"), ("}", r"\}"),
        ("~", r"\textasciitilde{}"), ("^", r"\textasciicircum{}"),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def _build_latex(title: str, author: str, sections: list, footer_str: str = "KI generierter Inhalt") -> str:
    esc = _latex_escape

    lines = [
        r"\documentclass[11pt,a4paper]{article}",
        r"\usepackage[a4paper,top=25mm,bottom=25mm,left=25mm,right=25mm]{geometry}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage[T1]{fontenc}",
        r"\usepackage{amsmath,amssymb}",
        r"\usepackage{booktabs}",
        r"\usepackage{tabularx}",
        r"\usepackage{parskip}",
        r"\usepackage{xcolor}",
        r"\usepackage{lmodern}",
        r"\usepackage[colorlinks=true,linkcolor=blue!60!black,urlcolor=blue!60!black]{hyperref}",
        r"\usepackage{fancyhdr}",
        r"\pagestyle{fancy}",
        r"\fancyhf{}",
        r"\rhead{\small\textcolor{gray}{AI_Framework_Thomas — Ingenieurbericht}}",
        r"\cfoot{\thepage}",
        rf"\lfoot{{\small\textcolor{{gray}}{{{_latex_escape(footer_str)}}}}}",
        r"\setlength{\headheight}{14pt}",
        "",
        rf"\title{{\textbf{{{esc(title)}}}}}",
        rf"\author{{{esc(author)}}}" if author else r"\date{}",
        r"\date{\today}",
        "",
        r"\begin{document}",
        r"\maketitle",
        r"\tableofcontents",
        r"\newpage",
        "",
    ]

    for sec in sections:
        heading = sec.get("heading", "")
        content = sec.get("content", "")
        equations = sec.get("equations", [])
        table = sec.get("table")
        subsections = sec.get("subsections", [])

        if heading:
            lines.append(rf"\section{{{esc(heading)}}}")

        if content:
            # Preserve line breaks
            for para in content.split("\n\n"):
                para = para.strip()
                if para:
                    lines.append(esc(para) + "\n")

        for eq in equations:
            lines += [r"\begin{equation}", eq, r"\end{equation}", ""]

        if table:
            headers = table.get("headers", [])
            rows = table.get("rows", [])
            if headers:
                col_spec = "l" + "r" * (len(headers) - 1)
                lines += [
                    r"\begin{center}",
                    rf"\begin{{tabular}}{{{col_spec}}}",
                    r"\toprule",
                    " & ".join(r"\textbf{" + esc(str(h)) + "}" for h in headers) + r" \\",
                    r"\midrule",
                ]
                for row in rows:
                    lines.append(" & ".join(esc(str(c)) for c in row) + r" \\")
                lines += [r"\bottomrule", r"\end{tabular}", r"\end{center}", ""]

        for sub in subsections:
            sub_heading = sub.get("heading", "")
            sub_content = sub.get("content", "")
            if sub_heading:
                lines.append(rf"\subsection{{{esc(sub_heading)}}}")
            if sub_content:
                for para in sub_content.split("\n\n"):
                    para = para.strip()
                    if para:
                        lines.append(esc(para) + "\n")

    lines.append(r"\end{document}")
    return "\n".join(lines)


def generate_report(title: str, author: str = "", sections: list | None = None, profile: dict | None = None) -> str:
    """
    Creates an engineering PDF report via LaTeX (falls back to DOCX).
    Returns a user-facing message with the download link.
    """
    if sections is None:
        sections = []

    report_id = uuid.uuid4().hex[:10]
    profile = profile or {}
    footer_parts = []
    name = " ".join(filter(None, [profile.get("first_name"), profile.get("last_name")])).strip()
    company = profile.get("company", "").strip()
    if name:
        footer_parts.append(name)
    if company:
        footer_parts.append(company)
    footer_parts.append("KI generierter Inhalt")
    footer_str = " · ".join(footer_parts)

    # Try LaTeX → PDF
    try:
        return _compile_pdf(title, author, sections, report_id, footer_str)
    except Exception as pdf_err:
        pass

    # Fallback: DOCX
    try:
        return _create_docx(title, author, sections, report_id, footer_str)
    except Exception as docx_err:
        return f"Report-Erstellung fehlgeschlagen: LaTeX: {pdf_err}, DOCX: {docx_err}"


def _compile_pdf(title, author, sections, report_id, footer_str="KI generierter Inhalt") -> str:
    latex_src = _build_latex(title, author, sections, footer_str)
    with tempfile.TemporaryDirectory() as tmp:
        tex_file = Path(tmp) / "report.tex"
        tex_file.write_text(latex_src, encoding="utf-8")
        result = subprocess.run(
            ["pdflatex", "-interaction=nonstopmode", "report.tex"],
            cwd=tmp, capture_output=True, text=True, timeout=30
        )
        # Run twice for ToC
        subprocess.run(
            ["pdflatex", "-interaction=nonstopmode", "report.tex"],
            cwd=tmp, capture_output=True, timeout=30
        )
        pdf_src = Path(tmp) / "report.pdf"
        if not pdf_src.exists():
            raise RuntimeError(f"pdflatex failed:\n{result.stderr[-500:]}")
        dest = REPORTS_DIR / f"{report_id}.pdf"
        dest.write_bytes(pdf_src.read_bytes())

    return (
        f"PDF-Bericht erstellt: **{title}**\n\n"
        f"[⬇ Herunterladen](/api/downloads/{report_id}.pdf)"
    )


def _create_docx(title, author, sections, report_id, footer_str="KI generierter Inhalt") -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(title, 0)
    if author:
        p = doc.add_paragraph(author)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for sec in sections:
        heading = sec.get("heading", "")
        content = sec.get("content", "")
        table = sec.get("table")

        if heading:
            doc.add_heading(heading, 1)
        if content:
            doc.add_paragraph(content)

        if table:
            headers = table.get("headers", [])
            rows = table.get("rows", [])
            if headers:
                t = doc.add_table(rows=1 + len(rows), cols=len(headers))
                t.style = "Light Grid Accent 1"
                for i, h in enumerate(headers):
                    t.rows[0].cells[i].text = str(h)
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        t.rows[ri + 1].cells[ci].text = str(cell)

    # Footer
    section = doc.sections[0]
    footer_para = section.footer.paragraphs[0]
    footer_para.text = footer_str
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer_para.runs:
        footer_para.runs[0].font.size = Pt(8)
        footer_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    dest = REPORTS_DIR / f"{report_id}.docx"
    doc.save(str(dest))
    return (
        f"DOCX-Bericht erstellt: **{title}**\n\n"
        f"[⬇ Herunterladen](/api/downloads/{report_id}.docx)"
    )
