"""
Datei-Extraktion für AI_Framework_Thomas — wandelt hochgeladene Dokumente in Text um.

Unterstützt PDF (pypdf), DOCX (python-docx), XLSX/CSV (openpyxl) sowie
Klartextformate. Bilder werden nicht hier verarbeitet, sondern im Backend
Base64-kodiert direkt an multimodale Modelle übergeben.

Einstiegspunkt: :func:`extract` — wählt anhand der Dateiendung den passenden
Parser und gibt den extrahierten Text zurück.
"""
from pathlib import Path


def read_table(fp: Path, sheet=None, header_row: int = 0, max_rows=None) -> dict:
    """Liest eine XLSX/CSV-Datei **strukturiert** ein (im Gegensatz zu :func:`_spreadsheet`,
    das nur Fließtext für RAG liefert und bei 200 Zeilen kappt).

    Rückgabe: ``{"sheets": [Blattnamen], "sheet": gewähltes Blatt, "headers": [Spalten],
    "rows": [[Zelle, …], …]}``. ``header_row`` ist der 0-basierte Index der Kopfzeile;
    alle Zeilen darüber werden übersprungen, die Kopfzeile selbst liefert die Spaltennamen.
    ``max_rows`` begrenzt die zurückgegebenen Datenzeilen (None = unbegrenzt)."""
    suffix = fp.suffix.lower()
    if suffix == ".csv":
        return _read_csv_table(fp, header_row, max_rows)
    return _read_xlsx_table(fp, sheet, header_row, max_rows)


def _cells_to_strs(values) -> list:
    return [("" if v is None else str(v)).strip() for v in values]


def _split_header_rows(all_rows: list, header_row: int, max_rows):
    """Teilt rohe Zeilen in (headers, data_rows) anhand der Kopfzeile. Datenzeilen,
    die komplett leer sind, werden verworfen; Zeilen werden auf die Header-Breite
    aufgefüllt/gekürzt, damit das Raster rechteckig ist."""
    if header_row < 0:
        header_row = 0
    if header_row >= len(all_rows):
        return [], []
    headers = _cells_to_strs(all_rows[header_row])
    # Leere Endspalten der Kopfzeile abschneiden
    while headers and headers[-1] == "":
        headers.pop()
    if not headers:
        return [], []
    width = len(headers)
    data = []
    for raw in all_rows[header_row + 1:]:
        cells = _cells_to_strs(raw)
        if not any(cells):
            continue
        cells = (cells + [""] * width)[:width]
        data.append(cells)
        if max_rows is not None and len(data) >= max_rows:
            break
    return headers, data


def _read_xlsx_table(fp: Path, sheet, header_row: int, max_rows) -> dict:
    try:
        import openpyxl
    except ImportError:
        return {"sheets": [], "sheet": "", "headers": [], "rows": [], "error": "openpyxl nicht installiert"}
    try:
        wb = openpyxl.load_workbook(str(fp), read_only=True, data_only=True)
    except Exception as e:
        return {"sheets": [], "sheet": "", "headers": [], "rows": [], "error": f"Tabellen-Fehler: {e}"}
    sheets = list(wb.sheetnames)
    name = sheet if sheet in sheets else (sheets[0] if sheets else "")
    if not name:
        return {"sheets": sheets, "sheet": "", "headers": [], "rows": []}
    ws = wb[name]
    all_rows = [list(r) for r in ws.iter_rows(values_only=True)]
    headers, data = _split_header_rows(all_rows, header_row, max_rows)
    return {"sheets": sheets, "sheet": name, "headers": headers, "rows": data}


def _read_csv_table(fp: Path, header_row: int, max_rows) -> dict:
    import csv, io
    text = fp.read_text(errors="replace")
    # Trennzeichen heuristisch erkennen (; oder , oder Tab)
    sample = text[:4000]
    delim = ";" if sample.count(";") >= sample.count(",") else ","
    if sample.count("\t") > sample.count(delim):
        delim = "\t"
    all_rows = [row for row in csv.reader(io.StringIO(text), delimiter=delim)]
    headers, data = _split_header_rows(all_rows, header_row, max_rows)
    return {"sheets": ["CSV"], "sheet": "CSV", "headers": headers, "rows": data}


def extract(fp: Path) -> str:
    suffix = fp.suffix.lower()

    if suffix == ".pdf":
        return _pdf(fp)
    elif suffix in (".docx", ".doc"):
        return _docx(fp)
    elif suffix in (".xlsx", ".xls", ".csv"):
        return _spreadsheet(fp)
    elif suffix in (".txt", ".md", ".py", ".js", ".json", ".yaml", ".yml", ".html", ".css"):
        return fp.read_text(errors="replace")
    elif suffix in (".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp"):
        return f"[Bild: {fp.name} — wird direkt an Bildmodell gesendet]"
    else:
        try:
            return fp.read_text(errors="replace")
        except Exception as e:
            return f"[Kann Datei nicht lesen: {e}]"


def _pdf(fp: Path) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(fp))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"--- Seite {i+1} ---\n{text}")
        return "\n\n".join(pages) if pages else "[PDF ohne lesbaren Text (möglicherweise gescannt)]"
    except ImportError:
        return "[pypdf nicht installiert]"
    except Exception as e:
        return f"[PDF-Fehler: {e}]"


def _docx(fp: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(fp))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append(" | ".join(cell.text for cell in row.cells))
            tables.append("\n".join(rows))
        content = "\n\n".join(paragraphs)
        if tables:
            content += "\n\n[Tabellen]\n" + "\n\n".join(tables)
        return content or "[Leeres Dokument]"
    except ImportError:
        return "[python-docx nicht installiert]"
    except Exception as e:
        return f"[DOCX-Fehler: {e}]"


def _spreadsheet(fp: Path) -> str:
    suffix = fp.suffix.lower()
    try:
        if suffix == ".csv":
            return fp.read_text(errors="replace")

        import openpyxl
        wb = openpyxl.load_workbook(str(fp), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True, max_row=200):
                row_str = " | ".join(str(c) if c is not None else "" for c in row)
                if row_str.strip(" |"):
                    rows.append(row_str)
            if rows:
                parts.append(f"[Blatt: {sheet_name}]\n" + "\n".join(rows))
        return "\n\n".join(parts) or "[Leere Tabelle]"
    except ImportError:
        return "[openpyxl nicht installiert]"
    except Exception as e:
        return f"[Tabellen-Fehler: {e}]"
