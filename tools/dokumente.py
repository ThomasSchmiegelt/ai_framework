"""Dokumentenerzeugung für die Tabs „Rechnungen" und „Arbeitszeugnisse".

Reine Logik ohne FastAPI/DB-Bezug (analog ``tools/mailstore.py``):
- **Rechnungen:** deterministische Betragsberechnung mit ``Decimal`` (Netto/USt/
  Brutto, §14-UStG-Pflichtangaben), Aufbau als Markdown (für den vorhandenen
  PDF-Renderer ``tools.export.to_pdf``) sowie ein eigener DOCX-Bauer mit echter
  Positionstabelle (``python-docx``).
- **Arbeitszeugnisse:** Prompt-/Struktur-Helfer; der eigentliche Text kommt vom
  LLM (in ``main.py`` aufgerufen), das Rendern übernehmen die generischen
  Exporter ``tools.export.to_pdf`` / ``to_docx`` (reiner Fließtext).

Die Beträge werden **nie** vom LLM berechnet — nur hier, damit Rechnungen
rechnerisch korrekt sind.
"""

from __future__ import annotations

import tempfile
from datetime import date, datetime, timedelta
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Optional

# ── Geldbeträge ───────────────────────────────────────────────────────────────

_CENT = Decimal("0.01")


def _money(value) -> Decimal:
    """Robuste Umwandlung nach Decimal mit 2 Nachkommastellen (kaufm. Rundung).
    Akzeptiert deutsche (``1.234,56``) wie englische (``1234.56``) Schreibweise."""
    if isinstance(value, Decimal):
        d = value
    elif isinstance(value, (int, float)):
        d = Decimal(str(value))
    else:
        s = str(value or "0").strip().replace("€", "").replace(" ", "")
        if not s:
            s = "0"
        # deutsche Notation: Punkt = Tausender, Komma = Dezimal
        if "," in s and ("." not in s or s.rfind(",") > s.rfind(".")):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
        try:
            d = Decimal(s)
        except Exception:
            d = Decimal("0")
    return d.quantize(_CENT, rounding=ROUND_HALF_UP)


def _qty(value) -> Decimal:
    """Menge als Decimal (bis 3 Nachkommastellen, für Stunden/Teilmengen)."""
    if isinstance(value, (int, float, Decimal)):
        d = Decimal(str(value))
    else:
        s = str(value or "1").strip().replace(" ", "")
        if "," in s and "." not in s:
            s = s.replace(",", ".")
        try:
            d = Decimal(s)
        except Exception:
            d = Decimal("1")
    return d.quantize(Decimal("0.001"), rounding=ROUND_HALF_UP)


def fmt_eur(d: Decimal) -> str:
    """Decimal → ``1.234,56 €`` (deutsche Formatierung)."""
    d = Decimal(d).quantize(_CENT, rounding=ROUND_HALF_UP)
    neg = d < 0
    s = f"{abs(d):,.2f}"                    # 1,234.56
    s = s.replace(",", "§").replace(".", ",").replace("§", ".")  # → 1.234,56
    return ("-" if neg else "") + s + " €"


def fmt_qty(d: Decimal) -> str:
    """Menge ohne überflüssige Nullen (``3`` statt ``3,000``; ``1,5`` bleibt)."""
    d = Decimal(d)
    s = format(d.normalize(), "f")
    if "." in s:
        s = s.rstrip("0").rstrip(".")
    return s.replace(".", ",")


def _fmt_date(value) -> str:
    """ISO-Datum oder Freitext → ``TT.MM.JJJJ`` (Freitext bleibt unverändert)."""
    if not value:
        return ""
    if isinstance(value, (date, datetime)):
        return value.strftime("%d.%m.%Y")
    s = str(value).strip()
    for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%Y/%m/%d"):
        try:
            return datetime.strptime(s, fmt).strftime("%d.%m.%Y")
        except ValueError:
            continue
    return s


# Standard-Leistungskategorien für die Vorgangs-Zerlegung (editierbar im Frontend).
RECHNUNG_KATEGORIEN = [
    "Recherche", "Planung", "Konstruktion", "Beschaffung", "Fremdleistungen",
    "Fertigung/Montage", "Inbetriebnahme", "Dokumentation", "Projektmanagement",
    "Programmierung", "Prüfung/Qualitätssicherung", "Reisezeit",
]


# ── Rechnungsberechnung ───────────────────────────────────────────────────────

def compute_invoice(inv: dict) -> dict:
    """Ergänzt eine Rechnung um berechnete Beträge.

    Erwartet ``positionen: [{menge, einheit, beschreibung, einzelpreis}]``,
    ``ust_satz`` (Prozent, z. B. 19) und ``kleinunternehmer`` (bool).
    Gibt eine Kopie mit ``positionen[*].betrag``, ``summe_netto``, ``ust_betrag``,
    ``summe_brutto`` (alles ``Decimal``) zurück."""
    out = dict(inv)
    kleinunternehmer = bool(inv.get("kleinunternehmer"))
    satz = Decimal("0") if kleinunternehmer else Decimal(str(inv.get("ust_satz", 19) or 0))

    positionen = []
    netto = Decimal("0.00")
    for p in (inv.get("positionen") or []):
        menge = _qty(p.get("menge", 1))
        einzel = _money(p.get("einzelpreis", 0))
        betrag = _money(menge * einzel)
        netto += betrag
        positionen.append({
            "menge": menge,
            "einheit": (p.get("einheit") or "").strip(),
            "beschreibung": (p.get("beschreibung") or "").strip(),
            "einzelpreis": einzel,
            "betrag": betrag,
        })

    netto = _money(netto)
    ust_betrag = _money(netto * satz / Decimal("100"))
    brutto = _money(netto + ust_betrag)

    out["positionen"] = positionen
    out["ust_satz"] = satz
    out["kleinunternehmer"] = kleinunternehmer
    out["summe_netto"] = netto
    out["ust_betrag"] = ust_betrag
    out["summe_brutto"] = brutto
    return out


def _basis_date(inv: dict) -> "date":
    """Belegdatum als ``date`` (heute, falls unlesbar)."""
    basis = inv.get("datum")
    if basis:
        for fmt in ("%Y-%m-%d", "%d.%m.%Y"):
            try:
                return datetime.strptime(str(basis).strip(), fmt).date()
            except ValueError:
                continue
    return date.today()


def _faellig(inv: dict) -> str:
    """Fälligkeitsdatum aus Rechnungsdatum + Zahlungsziel (Tage)."""
    tage = int(inv.get("zahlungsziel_tage", 14) or 0)
    return (_basis_date(inv) + timedelta(days=tage)).strftime("%d.%m.%Y")


def _gueltig_bis(inv: dict) -> str:
    """Gültigkeitsdatum eines Angebots: explizites Feld ``gueltig_bis`` oder
    Angebotsdatum + ``gueltig_tage`` (Standard 30)."""
    explizit = inv.get("gueltig_bis")
    if explizit:
        return _fmt_date(explizit)
    tage = int(inv.get("gueltig_tage", 30) or 0)
    return (_basis_date(inv) + timedelta(days=tage)).strftime("%d.%m.%Y")


# ── Rechnung → Markdown (für tools.export.to_pdf) ─────────────────────────────

def invoice_markdown(inv: dict, profile: dict, typ: str = "rechnung") -> str:
    """Berechnete oder unberechnete Rechnung/Angebot → Markdown-Fließtext.

    ``typ`` = ``"rechnung"`` (Standard) oder ``"angebot"`` — steuert Titel,
    Meta-Beschriftungen, Einleitung und den Abschlussblock (Zahlung vs.
    Gültigkeit). Positions-, USt- und Summenblock sind identisch.
    Nutzt eine Markdown-Tabelle für die Positionen (der PDF-Renderer
    ``_pdf_blocks_from_markdown`` versteht ``|``-Tabellen)."""
    inv = compute_invoice(inv)   # idempotent — stellt betrag/Summen sicher
    is_angebot = (typ == "angebot")
    titel = "Angebot" if is_angebot else "Rechnung"
    datum_label = "Angebotsdatum" if is_angebot else "Rechnungsdatum"
    nummer_label = "Angebotsnummer" if is_angebot else "Rechnungsnummer"
    betrag_label = "Angebotssumme" if is_angebot else "Rechnungsbetrag"
    p = profile or {}
    kunde = inv.get("kunde") or {}
    lines: list[str] = []

    # Absenderzeile (klein, einzeilig) + Empfängerblock
    absender_zeile = " · ".join(
        x for x in [p.get("firma"), p.get("strasse"), p.get("plz_ort")] if x)
    if absender_zeile:
        lines.append(f"*{absender_zeile}*")
        lines.append("")

    for feld in ["name", "zusatz", "strasse", "plz_ort", "land"]:
        val = (kunde.get(feld) or "").strip()
        if val:
            lines.append(val)
    lines.append("")

    lines.append(f"# {titel} {inv.get('nummer', '')}".rstrip())
    meta = []
    if inv.get("datum"):
        meta.append(f"**{datum_label}:** {_fmt_date(inv.get('datum'))}")
    if inv.get("nummer"):
        meta.append(f"**{nummer_label}:** {inv.get('nummer')}")
    if is_angebot:
        meta.append(f"**Gültig bis:** {_gueltig_bis(inv)}")
    if inv.get("leistungsdatum"):
        label = "Ausführungszeitraum" if is_angebot else "Leistungszeitraum"
        meta.append(f"**{label}:** {inv.get('leistungsdatum')}")
    if kunde.get("kundennummer"):
        meta.append(f"**Kundennummer:** {kunde.get('kundennummer')}")
    if meta:
        lines.append("  \n".join(meta))
        lines.append("")

    einleitung = str(inv.get("einleitung") or "").strip()
    if not einleitung and is_angebot:
        einleitung = ("Vielen Dank für Ihre Anfrage. Gerne unterbreiten wir Ihnen "
                      "folgendes Angebot:")
    if einleitung:
        lines.append(einleitung)
        lines.append("")

    # Positionstabelle
    lines.append("| Pos. | Menge | Beschreibung | Einzelpreis | Betrag |")
    lines.append("|---|---|---|---|---|")
    for i, pos in enumerate(inv["positionen"], 1):
        menge = f"{fmt_qty(pos['menge'])} {pos.get('einheit', '')}".strip()
        lines.append(
            f"| {i} | {menge} | {pos['beschreibung']} | "
            f"{fmt_eur(pos['einzelpreis'])} | {fmt_eur(pos['betrag'])} |")
    lines.append("")

    # Summenblock
    lines.append(f"**Summe netto:** {fmt_eur(inv['summe_netto'])}")
    if inv.get("kleinunternehmer"):
        lines.append("")
        lines.append("Gemäß § 19 UStG wird keine Umsatzsteuer berechnet.")
        lines.append("")
        lines.append(f"**{betrag_label}:** {fmt_eur(inv['summe_brutto'])}")
    else:
        satz = inv["ust_satz"]
        satz_str = fmt_qty(satz)
        lines.append(f"**zzgl. {satz_str} % USt:** {fmt_eur(inv['ust_betrag'])}")
        lines.append(f"**{betrag_label} (brutto):** {fmt_eur(inv['summe_brutto'])}")
    lines.append("")

    # Abschlussblock: Angebot = Gültigkeit, Rechnung = Zahlungsbedingungen
    if is_angebot:
        lines.append(f"Dieses Angebot ist freibleibend und gültig bis zum "
                     f"{_gueltig_bis(inv)}. Wir freuen uns auf Ihren Auftrag.")
    else:
        faellig = _faellig(inv)
        zahl = (f"Bitte überweisen Sie den Rechnungsbetrag von "
                f"{fmt_eur(inv['summe_brutto'])} bis zum {faellig}")
        bank_bits = [x for x in [
            (f"IBAN {p.get('iban')}" if p.get("iban") else ""),
            (f"BIC {p.get('bic')}" if p.get("bic") else ""),
            (f"{p.get('bank')}" if p.get("bank") else ""),
        ] if x]
        if bank_bits:
            zahl += " auf folgendes Konto: " + ", ".join(bank_bits) + "."
        else:
            zahl += "."
        lines.append(zahl)
    if inv.get("hinweis"):
        lines.append("")
        lines.append(str(inv["hinweis"]).strip())

    # Pflichtangaben-Fußblock (§14 UStG)
    lines.append("")
    lines.append("---")
    fuss = []
    if p.get("firma"):
        fuss.append(p["firma"])
    if p.get("inhaber"):
        fuss.append(p["inhaber"])
    anschrift = " · ".join(x for x in [p.get("strasse"), p.get("plz_ort")] if x)
    if anschrift:
        fuss.append(anschrift)
    kontakt = " · ".join(x for x in [
        (f"Tel. {p.get('telefon')}" if p.get("telefon") else ""),
        p.get("email") or "",
    ] if x)
    if kontakt:
        fuss.append(kontakt)
    steuer = " · ".join(x for x in [
        (f"USt-IdNr. {p.get('ust_id')}" if p.get("ust_id") else ""),
        (f"Steuernummer {p.get('steuernummer')}" if p.get("steuernummer") else ""),
    ] if x)
    if steuer:
        fuss.append(steuer)
    for f in fuss:
        lines.append(f"*{f}*  ")

    return "\n".join(lines)


# ── Rechnung → DOCX (eigene Positionstabelle) ─────────────────────────────────

def invoice_docx(inv: dict, profile: dict, typ: str = "rechnung") -> Path:
    """Rechnung/Angebot als DOCX mit echter Positionstabelle. Gibt den Pfad zurück.

    ``typ`` = ``"rechnung"`` (Standard) oder ``"angebot"`` (siehe
    ``invoice_markdown``)."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    inv = compute_invoice(inv)   # idempotent — stellt betrag/Summen sicher
    is_angebot = (typ == "angebot")
    titel = "Angebot" if is_angebot else "Rechnung"
    datum_label = "Angebotsdatum" if is_angebot else "Rechnungsdatum"
    nummer_label = "Angebotsnummer" if is_angebot else "Rechnungsnummer"
    betrag_label = "Angebotssumme" if is_angebot else "Rechnungsbetrag"
    p = profile or {}
    kunde = inv.get("kunde") or {}

    doc = Document()

    def _small(text, italic=True, gray=True):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(8)
        run.font.italic = italic
        if gray:
            run.font.color.rgb = RGBColor(0x6C, 0x6F, 0x76)
        return para

    absender_zeile = " · ".join(
        x for x in [p.get("firma"), p.get("strasse"), p.get("plz_ort")] if x)
    if absender_zeile:
        _small(absender_zeile)

    # Empfänger
    for feld in ["name", "zusatz", "strasse", "plz_ort", "land"]:
        val = (kunde.get(feld) or "").strip()
        if val:
            doc.add_paragraph(val)

    doc.add_paragraph()
    doc.add_heading(f"{titel} {inv.get('nummer', '')}".rstrip(), level=1)

    def _meta(label, value):
        if not value:
            return
        para = doc.add_paragraph()
        r = para.add_run(f"{label}: ")
        r.bold = True
        para.add_run(str(value))

    _meta(nummer_label, inv.get("nummer"))
    _meta(datum_label, _fmt_date(inv.get("datum")))
    if is_angebot:
        _meta("Gültig bis", _gueltig_bis(inv))
    _meta("Ausführungszeitraum" if is_angebot else "Leistungszeitraum",
          inv.get("leistungsdatum"))
    _meta("Kundennummer", kunde.get("kundennummer"))

    einleitung = str(inv.get("einleitung") or "").strip()
    if not einleitung and is_angebot:
        einleitung = ("Vielen Dank für Ihre Anfrage. Gerne unterbreiten wir Ihnen "
                      "folgendes Angebot:")
    if einleitung:
        doc.add_paragraph()
        doc.add_paragraph(einleitung)

    # Positionstabelle
    doc.add_paragraph()
    headers = ["Pos.", "Menge", "Beschreibung", "Einzelpreis", "Betrag"]
    table = doc.add_table(rows=1 + len(inv["positionen"]), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True
    for ri, pos in enumerate(inv["positionen"], 1):
        menge = f"{fmt_qty(pos['menge'])} {pos.get('einheit', '')}".strip()
        vals = [str(ri), menge, pos["beschreibung"],
                fmt_eur(pos["einzelpreis"]), fmt_eur(pos["betrag"])]
        for ci, v in enumerate(vals):
            cell = table.cell(ri, ci)
            cell.text = v
            if ci in (3, 4):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Summenblock
    doc.add_paragraph()

    def _sum_line(label, value, bold=False):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = para.add_run(f"{label}  {value}")
        r.bold = bold

    _sum_line("Summe netto:", fmt_eur(inv["summe_netto"]))
    if inv.get("kleinunternehmer"):
        doc.add_paragraph("Gemäß § 19 UStG wird keine Umsatzsteuer berechnet.")
        _sum_line(f"{betrag_label}:", fmt_eur(inv["summe_brutto"]), bold=True)
    else:
        _sum_line(f"zzgl. {fmt_qty(inv['ust_satz'])} % USt:", fmt_eur(inv["ust_betrag"]))
        _sum_line(f"{betrag_label} (brutto):", fmt_eur(inv["summe_brutto"]), bold=True)

    # Abschlussblock: Angebot = Gültigkeit, Rechnung = Zahlungsbedingungen
    doc.add_paragraph()
    if is_angebot:
        doc.add_paragraph(
            f"Dieses Angebot ist freibleibend und gültig bis zum "
            f"{_gueltig_bis(inv)}. Wir freuen uns auf Ihren Auftrag.")
    else:
        faellig = _faellig(inv)
        zahl = (f"Bitte überweisen Sie den Rechnungsbetrag von "
                f"{fmt_eur(inv['summe_brutto'])} bis zum {faellig}")
        bank_bits = [x for x in [
            (f"IBAN {p.get('iban')}" if p.get("iban") else ""),
            (f"BIC {p.get('bic')}" if p.get("bic") else ""),
            (p.get("bank") or ""),
        ] if x]
        zahl += (" auf folgendes Konto: " + ", ".join(bank_bits) + ".") if bank_bits else "."
        doc.add_paragraph(zahl)
    if inv.get("hinweis"):
        doc.add_paragraph(str(inv["hinweis"]).strip())

    # Fußzeile: Pflichtangaben (§14 UStG)
    fuss_bits = [x for x in [
        p.get("firma"), p.get("inhaber"),
        " · ".join(y for y in [p.get("strasse"), p.get("plz_ort")] if y),
        " · ".join(y for y in [
            (f"Tel. {p.get('telefon')}" if p.get("telefon") else ""),
            p.get("email") or "",
        ] if y),
        " · ".join(y for y in [
            (f"USt-IdNr. {p.get('ust_id')}" if p.get("ust_id") else ""),
            (f"Steuernummer {p.get('steuernummer')}" if p.get("steuernummer") else ""),
        ] if y),
    ] if x]
    footer = doc.sections[0].footer
    fp_para = footer.paragraphs[0]
    fp_para.text = " | ".join(fuss_bits)
    fp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if fp_para.runs:
        fp_para.runs[0].font.size = Pt(7.5)
        fp_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    fp = Path(tempfile.mktemp(suffix=".docx"))
    doc.save(str(fp))
    return fp


# ── Plan → Angebotspositionen ─────────────────────────────────────────────────

def _resource_cost(r: dict) -> Decimal:
    """Kosten einer Plan-Ressource (identisch zur Frontend-Formel
    ``_resCost`` in planner.js: hours>0 ? qty*hours*rate : qty*rate)."""
    qty = _dec(r.get("qty"))
    hours = _dec(r.get("hours"))
    rate = _dec(r.get("rate"))
    return (qty * hours * rate) if hours > 0 else (qty * rate)


def _dec(value) -> Decimal:
    try:
        s = str(value if value not in (None, "") else 0).replace(",", ".")
        return Decimal(s)
    except Exception:
        return Decimal("0")


def _task_cost(t: dict) -> Decimal:
    """Gesamtkosten eines Vorgangs (Summe seiner Ressourcen)."""
    return sum((_resource_cost(r) for r in (t.get("resource_list") or [])),
               Decimal("0"))


def plan_to_positions(plan: dict) -> list[dict]:
    """Wandelt einen Planer-Plan in Angebotspositionen um.

    Sind **Bereiche** (``area`` / Spalte „Bereich") gepflegt, wird je Bereich
    **eine** Position mit der Summe der Vorgangskosten gebildet (Reihenfolge =
    erstes Vorkommen; Vorgänge ohne Bereich unter „Sonstiges").

    Hat **kein** Vorgang einen Bereich (typisch bei automatisch erzeugten
    Plänen), wäre eine einzige „Sonstiges"-Sammelposition wertlos — dann wird
    **je Vorgang eine Position** ausgegeben (Vorgangsname + dessen Kosten).
    Kostenlose Vorgänge (ohne Ressourcen, z. B. Meilensteine) werden dabei
    übersprungen; sind *alle* Vorgänge kostenlos, bleiben sie als 0-€-Zeilen
    erhalten, damit die Struktur sichtbar ist und Preise ergänzt werden können.

    Rückgabe passt zu ``compute_invoice``:
    ``[{menge, einheit, beschreibung, einzelpreis}]`` (Netto-Pauschale)."""
    tasks = list(plan.get("tasks") or [])
    has_areas = any(str(t.get("area") or "").strip() for t in tasks)

    if has_areas:
        order: list[str] = []
        sums: dict[str, Decimal] = {}
        for t in tasks:
            area = str(t.get("area") or "").strip() or "Sonstiges"
            if area not in sums:
                sums[area] = Decimal("0")
                order.append(area)
            sums[area] += _task_cost(t)
        return [{
            "menge": 1,
            "einheit": "pauschal",
            "beschreibung": area,
            "einzelpreis": float(sums[area].quantize(_CENT, rounding=ROUND_HALF_UP)),
        } for area in order]

    # Keine Bereiche gepflegt → je Vorgang eine Position.
    any_cost = any(_task_cost(t) > 0 for t in tasks)
    positionen = []
    for t in tasks:
        kosten = _task_cost(t)
        if any_cost and kosten <= 0:
            continue  # Meilensteine/kostenlose Vorgänge auslassen
        name = str(t.get("name") or "").strip() or "Vorgang"
        positionen.append({
            "menge": 1,
            "einheit": "pauschal",
            "beschreibung": name,
            "einzelpreis": float(kosten.quantize(_CENT, rounding=ROUND_HALF_UP)),
        })
    return positionen


# ── Arbeitszeugnis: Prompt-/Struktur-Helfer ───────────────────────────────────

# Gesamtnote → codierte Zeugnissprache (üblicher „Geheimcode")
NOTEN = {
    "1": {
        "label": "sehr gut",
        "leistung": "stets zu unserer vollsten Zufriedenheit",
        "verhalten": "stets vorbildlich",
        "schluss": "Wir bedauern sein/ihr Ausscheiden sehr und danken für die stets "
                   "hervorragende Zusammenarbeit. Für die Zukunft wünschen wir alles "
                   "Gute und weiterhin viel Erfolg.",
    },
    "2": {
        "label": "gut",
        "leistung": "stets zu unserer vollen Zufriedenheit",
        "verhalten": "stets einwandfrei",
        "schluss": "Wir bedauern sein/ihr Ausscheiden und danken für die gute "
                   "Zusammenarbeit. Für die Zukunft wünschen wir alles Gute und viel Erfolg.",
    },
    "3": {
        "label": "befriedigend",
        "leistung": "zu unserer vollen Zufriedenheit",
        "verhalten": "einwandfrei",
        "schluss": "Wir danken für die Zusammenarbeit und wünschen für die Zukunft "
                   "alles Gute.",
    },
    "4": {
        "label": "ausreichend",
        "leistung": "zu unserer Zufriedenheit",
        "verhalten": "insgesamt ohne Beanstandungen",
        "schluss": "Wir wünschen für die Zukunft alles Gute.",
    },
    "5": {
        "label": "mangelhaft",
        "leistung": "im Großen und Ganzen zu unserer Zufriedenheit",
        "verhalten": "ohne wesentliche Beanstandungen",
        "schluss": "Wir wünschen für den weiteren Berufsweg alles Gute.",
    },
}


def note_key(value) -> str:
    """Normalisiert eine Note (1–5, ‚sehr gut'…) auf den Schlüssel ‚1'…‚5'."""
    s = str(value or "2").strip().lower()
    if s and s[0] in "12345":
        return s[0]
    for k, v in NOTEN.items():
        if v["label"] in s:
            return k
    return "2"


def zeugnis_system_prompt() -> str:
    return (
        "Du bist erfahrene:r Personalreferent:in und formulierst ein rechtssicheres, "
        "wohlwollendes qualifiziertes Arbeitszeugnis nach deutschem Arbeitsrecht "
        "(§ 109 GewO). Beachte strikt:\n"
        "- Verwende die übliche codierte Zeugnissprache passend zur vorgegebenen "
        "Gesamtnote (die Zufriedenheits-/Verhaltensformeln werden dir genannt und "
        "sind wortgetreu einzubauen).\n"
        "- Aufbau: (1) Einleitung mit Personalien, Position und Beschäftigungsdauer; "
        "(2) Aufgabenbeschreibung; (3) ggf. Fachwissen/Weiterbildung; "
        "(4) Leistungsbeurteilung (Arbeitsweise, Können, Erfolge) mit der "
        "vorgegebenen Leistungsformel; (5) Sozialverhalten gegenüber Vorgesetzten, "
        "Kolleg:innen und Kund:innen mit der vorgegebenen Verhaltensformel; "
        "(6) Schlussabsatz (Beendigungsgrund, Dank, Bedauern, Zukunftswünsche) "
        "passend zur Note.\n"
        "- Durchgehend positive, wohlwollende, aber wahrheitsgemäße Sprache; keine "
        "verdeckt-negativen Formulierungen über die Note hinaus.\n"
        "- Nutze durchgehend das angegebene Geschlecht/Pronomen und den vollen Namen.\n"
        "- Gib NUR den fertigen Zeugnistext als Fließtext (mit Absätzen) aus, keine "
        "Überschriften-Marker, keine Erklärungen, kein Vorwort."
    )


def zeugnis_user_prompt(meta: dict) -> str:
    nk = note_key(meta.get("note"))
    note = NOTEN[nk]
    geschlecht = (meta.get("geschlecht") or "divers").strip()
    aufgaben = meta.get("aufgaben") or ""
    if isinstance(aufgaben, list):
        aufgaben = "\n".join(f"- {a}" for a in aufgaben)
    zeilen = [
        f"Erstelle ein qualifiziertes Arbeitszeugnis mit Gesamtnote „{note['label']}\".",
        "",
        f"Arbeitgeber: {meta.get('arbeitgeber', '')}",
        f"Name der/des Beschäftigten: {meta.get('name', '')}",
        f"Geschlecht/Pronomen: {geschlecht}",
        f"Position/Funktion: {meta.get('position', '')}",
        f"Abteilung: {meta.get('abteilung', '')}",
        f"Beschäftigt von: {_fmt_date(meta.get('eintritt'))} bis {_fmt_date(meta.get('austritt'))}",
        f"Beendigungsgrund/Schlussformel-Wunsch: {meta.get('beendigung', 'einvernehmlich')}",
        "",
        "Aufgaben und Verantwortungsbereiche:",
        aufgaben or "(bitte branchenüblich zur Position ergänzen)",
    ]
    if meta.get("staerken"):
        zeilen += ["", f"Besondere Stärken/Erfolge: {meta['staerken']}"]
    if meta.get("fuehrung"):
        zeilen += ["", f"Führungsverantwortung: {meta['fuehrung']}"]
    zeilen += [
        "",
        "Wortgetreu einzubauende Formeln (Gesamtnote):",
        f"- Leistungsformel: „…{note['leistung']}.\"",
        f"- Verhaltensformel: Das Verhalten war {note['verhalten']}.",
        f"- Sinngemäßer Schlussabsatz: {note['schluss']}",
    ]
    return "\n".join(zeilen)


def zeugnis_markdown(text: str, meta: dict) -> str:
    """LLM-Zeugnistext → Markdown mit Titel/Ort-Datum-Zeile (für Exporter)."""
    md = ["# Arbeitszeugnis", ""]
    md.append(text.strip())
    md.append("")
    ort = (meta or {}).get("ort") or ""
    datum = _fmt_date((meta or {}).get("ausstellungsdatum")) or _fmt_date(date.today())
    ortdatum = ", ".join(x for x in [ort, datum] if x)
    if ortdatum:
        md += ["", ortdatum]
    if (meta or {}).get("unterzeichner"):
        md += ["", "", meta["unterzeichner"]]
    return "\n".join(md)
